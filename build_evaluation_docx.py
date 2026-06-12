from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "EVALUATION_REPORT.md"
OUTPUT = ROOT / "docs" / "EVALUATION_REPORT.docx"


def set_run_style(run, bold=False, italic=False, color=None, size=None, font="Arial"):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)


def add_inline_runs(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_style(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_style(run, font="Consolas", color=(15, 23, 42))
        else:
            run = paragraph.add_run(part)
            set_run_style(run)


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 16, (18, 107, 95)),
        ("Heading 2", 13, (29, 78, 216)),
        ("Heading 3", 11.5, (23, 32, 42)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title_style = styles["Title"]
    title_style.font.name = "Arial"
    title_style.font.bold = True
    title_style.font.size = Pt(22)
    title_style.font.color.rgb = RGBColor(16, 42, 67)


def add_header_footer(document):
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = "CS406.3 Evaluation Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_style(header.runs[0], bold=True, color=(99, 112, 132), size=8.5)

    footer = section.footer.paragraphs[0]
    footer.text = "Smart Employment Gateway for SLBFE"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_style(footer.runs[0], color=(99, 112, 132), size=8.5)


def add_table(document, table_lines):
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, text)
            for run in paragraph.runs:
                set_run_style(
                    run,
                    bold=r_idx == 0 or run.bold,
                    color=(255, 255, 255) if r_idx == 0 else (23, 32, 42),
                    size=9,
                )
            if r_idx == 0:
                shading = cell._tc.get_or_add_tcPr()
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "102A43")
                shading.append(shd)

    document.add_paragraph()


def add_code_block(document, code_lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("\n".join(code_lines))
    set_run_style(run, font="Consolas", size=9, color=(15, 23, 42))


def build_docx():
    document = Document()
    style_document(document)
    add_header_footer(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
      line = lines[i].rstrip()

      if not line:
          i += 1
          continue

      if line.startswith("# "):
          paragraph = document.add_paragraph(style="Title")
          paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
          add_inline_runs(paragraph, line[2:])
      elif line.startswith("## "):
          document.add_heading(line[3:], level=1)
      elif line.startswith("### "):
          document.add_heading(line[4:], level=2)
      elif line.startswith("```"):
          code_lines = []
          i += 1
          while i < len(lines) and not lines[i].startswith("```"):
              code_lines.append(lines[i])
              i += 1
          add_code_block(document, code_lines)
      elif line.startswith("|"):
          table_lines = [line]
          i += 1
          while i < len(lines) and lines[i].startswith("|"):
              table_lines.append(lines[i].rstrip())
              i += 1
          add_table(document, table_lines)
          continue
      elif line.startswith("- "):
          paragraph = document.add_paragraph(style="List Bullet")
          add_inline_runs(paragraph, line[2:])
      elif re.match(r"^\d+\.\s+", line):
          text = re.sub(r"^\d+\.\s+", "", line)
          paragraph = document.add_paragraph(style="List Number")
          add_inline_runs(paragraph, text)
      else:
          paragraph_lines = [line]
          i += 1
          while i < len(lines) and lines[i].strip() and not re.match(r"^(#|```|\||- |\d+\.\s+)", lines[i]):
              paragraph_lines.append(lines[i].strip())
              i += 1
          paragraph = document.add_paragraph()
          add_inline_runs(paragraph, " ".join(paragraph_lines))
          continue

      i += 1

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
